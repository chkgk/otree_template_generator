import docx
from ..token import get_token, merge_runs_with_same_style
from ..page import Page


class WordAdapter:
    def __init__(self, docx_file):
        self.pages = []
        self.player_fields = []
        self.group_fields = []
        self.docx_filename = docx_file
        self.docx_document = docx.Document(self.docx_filename)

    def parse(self):
        # we need to add some check to make sure that the very first token defines a page
        current_page = None
        for paragraph in self.docx_document.paragraphs:
            paragraph_content = []
            merged_runs = merge_runs_with_same_style(paragraph.runs)
            for run in merged_runs:
                s = get_token(run)
                if s.type == 'page':
                    current_page = Page(title=s.title, form_model=s.form_model)
                    self.pages.append(current_page)

                if current_page is None:
                    continue

                if s.type == 'input' and s.input_type != 'button':
                    current_page.form_fields.append(s)

                if s.type == 'output' and s.var_for_template:
                    current_page.vars_for_template.append(s.var_for_template)

                paragraph_content.append(s)

            p_prefix = p_postfix = ''
            if any([e.type == 'text' for e in paragraph_content]):
                p_prefix = '<p>'
                p_postfix = '</p>'

            rendered_elements = [e.render() for e in paragraph_content if e.render() is not None]
            if rendered_elements:
                content = "".join(rendered_elements)
                current_page.paragraphs.append(f"{p_prefix}{content}{p_postfix}")

        for page in self.pages:
            if page.form_model == 'player':
                self.player_fields.append(page.render_model_fields())

            if page.form_model == 'group':
                self.group_fields.append(page.render_model_fields())

